"""Document ingestion service — parse, chunk, embed, and store in Qdrant.

Supports: PDF, TXT, Markdown, DOCX.
Pipeline: Upload → Parse → Chunk → Embed → Qdrant upsert.
"""

import hashlib
import uuid
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Any, Optional

from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.config import settings
from app.core.logging import logger

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200


@dataclass
class IngestedDocument:
    """Metadata for an ingested document stored in Qdrant."""

    doc_id: str
    filename: str
    file_type: str
    chunk_count: int
    created_at: str
    user_id: str = ""
    file_size: int = 0
    metadata: dict[str, Any] = field(default_factory=dict)


def _parse_pdf(content: bytes) -> str:
    """Extract text from a PDF file."""
    import io

    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _parse_docx(content: bytes) -> str:
    """Extract text from a DOCX file."""
    import io

    from docx import Document

    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def _parse_text(content: bytes) -> str:
    """Extract text from a plain text / markdown file."""
    return content.decode("utf-8", errors="replace")


PARSERS = {
    ".pdf": _parse_pdf,
    ".txt": _parse_text,
    ".md": _parse_text,
    ".docx": _parse_docx,
}


def parse_document(filename: str, content: bytes) -> str:
    """Parse document content based on file extension.

    Args:
        filename: Original filename with extension.
        content: Raw file bytes.

    Returns:
        Extracted plain text.

    Raises:
        ValueError: If file type is not supported.
    """
    ext = _get_extension(filename)
    parser = PARSERS.get(ext)
    if parser is None:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}")
    return parser(content)


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Split text into chunks using recursive character splitter.

    Args:
        text: The full document text.
        chunk_size: Target chunk size in characters.
        chunk_overlap: Overlap between consecutive chunks.

    Returns:
        List of text chunks.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", ".", " ", ""],
    )
    return splitter.split_text(text)


async def ingest_document(
    filename: str,
    content: bytes,
    user_id: str = "",
    collection_name: Optional[str] = None,
    metadata: Optional[dict[str, Any]] = None,
) -> IngestedDocument:
    """Full ingestion pipeline: parse → chunk → embed → store in Qdrant.

    Args:
        filename: Original filename.
        content: Raw file bytes.
        user_id: ID of the uploading user.
        collection_name: Qdrant collection (default from settings).
        metadata: Extra metadata to attach to each chunk.

    Returns:
        IngestedDocument with ingestion result.
    """
    from langchain_openai import OpenAIEmbeddings
    from qdrant_client import AsyncQdrantClient
    from qdrant_client.models import Distance, PointStruct, VectorParams

    collection = collection_name or settings.QDRANT_COLLECTION_NAME
    doc_id = str(uuid.uuid4())
    ext = _get_extension(filename)
    extra_meta = metadata or {}

    logger.info(
        "document_ingestion_started",
        doc_id=doc_id,
        filename=filename,
        file_type=ext,
        file_size=len(content),
        user_id=user_id,
    )

    text = parse_document(filename, content)
    if not text.strip():
        raise ValueError("Document is empty or could not be parsed.")

    chunks = chunk_text(text)
    logger.info("document_chunked", doc_id=doc_id, chunk_count=len(chunks))

    embedding_model = settings.LONG_TERM_MEMORY_EMBEDDER_MODEL
    embedding_base_url = settings.LONG_TERM_MEMORY_EMBEDDER_BASE_URL or None
    embedding_dims = settings.LONG_TERM_MEMORY_EMBEDDER_DIMS

    embeddings_client = OpenAIEmbeddings(
        model=embedding_model,
        **({"openai_api_base": embedding_base_url} if embedding_base_url else {}),
    )
    vectors = await embeddings_client.aembed_documents(chunks)

    client = AsyncQdrantClient(
        host=settings.QDRANT_HOST,
        port=settings.QDRANT_PORT,
        api_key=settings.QDRANT_API_KEY or None,
    )

    try:
        collections = await client.get_collections()
        existing_names = [c.name for c in collections.collections]
        if collection not in existing_names:
            await client.create_collection(
                collection_name=collection,
                vectors_config=VectorParams(size=embedding_dims, distance=Distance.COSINE),
            )
            logger.info("qdrant_collection_created", collection=collection, dims=embedding_dims)

        now = datetime.now(timezone.utc).isoformat()
        points = []
        for i, (chunk, vector) in enumerate(zip(chunks, vectors, strict=True)):
            point_id = str(uuid.uuid4())
            payload = {
                "content": chunk,
                "source": filename,
                "doc_id": doc_id,
                "chunk_index": i,
                "user_id": user_id,
                "created_at": now,
                **extra_meta,
            }
            points.append(PointStruct(id=point_id, vector=vector, payload=payload))

        batch_size = 100
        for start in range(0, len(points), batch_size):
            batch = points[start : start + batch_size]
            await client.upsert(collection_name=collection, points=batch)

        logger.info(
            "document_ingestion_completed",
            doc_id=doc_id,
            filename=filename,
            chunk_count=len(chunks),
            collection=collection,
        )

        return IngestedDocument(
            doc_id=doc_id,
            filename=filename,
            file_type=ext,
            chunk_count=len(chunks),
            created_at=now,
            user_id=user_id,
            file_size=len(content),
            metadata=extra_meta,
        )
    finally:
        await client.close()


async def list_documents(
    user_id: str = "",
    collection_name: Optional[str] = None,
) -> list[dict[str, Any]]:
    """List all ingested documents by querying unique doc_ids from Qdrant.

    Args:
        user_id: Filter by user_id if provided.
        collection_name: Qdrant collection name.

    Returns:
        List of document metadata dicts.
    """
    from qdrant_client import AsyncQdrantClient
    from qdrant_client.models import FieldCondition, Filter, MatchValue, ScrollRequest

    collection = collection_name or settings.QDRANT_COLLECTION_NAME
    client = AsyncQdrantClient(
        host=settings.QDRANT_HOST,
        port=settings.QDRANT_PORT,
        api_key=settings.QDRANT_API_KEY or None,
    )

    try:
        collections = await client.get_collections()
        existing_names = [c.name for c in collections.collections]
        if collection not in existing_names:
            return []

        scroll_filter = None
        if user_id:
            scroll_filter = Filter(must=[FieldCondition(key="user_id", match=MatchValue(value=user_id))])

        seen_docs: dict[str, dict[str, Any]] = {}
        offset = None

        while True:
            results = await client.scroll(
                collection_name=collection,
                scroll_filter=scroll_filter,
                limit=100,
                offset=offset,
                with_payload=True,
                with_vectors=False,
            )
            points, next_offset = results

            for point in points:
                payload = point.payload or {}
                doc_id = payload.get("doc_id", "")
                if doc_id and doc_id not in seen_docs:
                    seen_docs[doc_id] = {
                        "doc_id": doc_id,
                        "filename": payload.get("source", ""),
                        "user_id": payload.get("user_id", ""),
                        "created_at": payload.get("created_at", ""),
                        "chunk_index": payload.get("chunk_index", 0),
                    }
                elif doc_id and doc_id in seen_docs:
                    existing = seen_docs[doc_id]
                    ci = payload.get("chunk_index", 0)
                    if ci > existing.get("chunk_index", 0):
                        existing["chunk_index"] = ci

            if next_offset is None:
                break
            offset = next_offset

        documents = []
        for doc_id, info in seen_docs.items():
            documents.append(
                {
                    "doc_id": doc_id,
                    "filename": info["filename"],
                    "user_id": info["user_id"],
                    "created_at": info["created_at"],
                    "chunk_count": info["chunk_index"] + 1,
                }
            )

        documents.sort(key=lambda d: d.get("created_at", ""), reverse=True)
        return documents
    finally:
        await client.close()


async def delete_document(
    doc_id: str,
    collection_name: Optional[str] = None,
) -> int:
    """Delete all chunks belonging to a document from Qdrant.

    Args:
        doc_id: The document ID to delete.
        collection_name: Qdrant collection name.

    Returns:
        Number of points deleted (approximate).
    """
    from qdrant_client import AsyncQdrantClient
    from qdrant_client.models import FieldCondition, Filter, MatchValue

    collection = collection_name or settings.QDRANT_COLLECTION_NAME
    client = AsyncQdrantClient(
        host=settings.QDRANT_HOST,
        port=settings.QDRANT_PORT,
        api_key=settings.QDRANT_API_KEY or None,
    )

    try:
        await client.delete(
            collection_name=collection,
            points_selector=Filter(must=[FieldCondition(key="doc_id", match=MatchValue(value=doc_id))]),
        )
        logger.info("document_deleted", doc_id=doc_id, collection=collection)
        return 1
    finally:
        await client.close()


def _get_extension(filename: str) -> str:
    """Extract lowercase file extension from filename."""
    import os

    _, ext = os.path.splitext(filename.lower())
    return ext
